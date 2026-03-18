"""Word document generation service using python-docx."""

from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from src.models.export_models import ExportData
from src.models.ai import StyleContentType
from src.models.vocabulary_audit import CONFIDENCE_THRESHOLDS


# Government of Canada primary color
GC_PRIMARY = RGBColor(0x26, 0x37, 0x4a)

# Confidence indicator colors (matching PDF CSS)
CONFIDENCE_COLORS = {
    "high": RGBColor(0x4c, 0xaf, 0x50),    # Green #4caf50
    "medium": RGBColor(0xff, 0x98, 0x00),  # Amber #ff9800
    "low": RGBColor(0xf4, 0x43, 0x36),     # Red #f44336
}

# AI disclosure blue (matching PDF CSS)
AI_DISCLOSURE_BLUE = RGBColor(0x19, 0x76, 0xd2)

# Source tag mapping (matches pdf_generator.py SOURCE_TAG_MAP)
SOURCE_TAG_MAP = {
    "Main Duties": "[NOC]",
    "Work Activities": "[OaSIS]",
    "Abilities": "[OaSIS]",
    "Knowledge": "[OaSIS]",
    "Skills": "[OaSIS]",
    "Effort": "[OaSIS]",
    "Responsibility": "[OaSIS]",
    "Core Competencies": "[GC]",
}


def get_source_tag(source_attribute: str, data_source: str = None) -> str:
    """Return the bracketed source tag for a statement."""
    if data_source == "jobforge":
        return "[JobForge]"
    return SOURCE_TAG_MAP.get(source_attribute, "[NOC]")


def _get_confidence_level(score: float) -> str:
    """Map confidence score to level string for color lookup."""
    if score >= CONFIDENCE_THRESHOLDS["high"]:
        return "high"
    elif score >= CONFIDENCE_THRESHOLDS["medium"]:
        return "medium"
    return "low"


def _add_hyperlink(paragraph, url, text, color='1565C0', underline=True):
    """Add a hyperlink to a paragraph.

    Args:
        paragraph: docx paragraph object
        url: The URL to link to
        text: Display text for the link
        color: Hex color code (without #)
        underline: Whether to underline the link

    Returns:
        The hyperlink run
    """
    # Create relationship (rId)
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run within hyperlink
    run_element = OxmlElement('w:r')

    # Run properties
    rPr = OxmlElement('w:rPr')

    # Color
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    run_element.append(rPr)

    # Text element
    text_element = OxmlElement('w:t')
    text_element.text = text
    run_element.append(text_element)

    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)

    return hyperlink


def _add_classification_section(doc, data):
    """Add classification section to document.

    Args:
        doc: Document object
        data: ExportData with classification_result
    """
    if not data.include_classification or not data.classification_result:
        return

    # Page break before classification
    doc.add_page_break()

    # Main heading
    heading = doc.add_heading('Classification Step 1: Occupational Group Allocation', 1)
    for run in heading.runs:
        run.font.color.rgb = GC_PRIMARY

    # Status table
    _add_key_value_table(doc, [
        ('Status', data.classification_result.get('status', '')),
        ('Match Context', data.classification_result.get('match_context', ''))
    ])

    # Recommendations heading
    rec_heading = doc.add_heading('Recommendations (Ranked by Confidence)', 2)
    for run in rec_heading.runs:
        run.font.color.rgb = GC_PRIMARY

    # Each recommendation
    for rec in data.classification_result.get('recommendations', []):
        # Recommendation heading
        rec_title = f"{rec['group_code']}: {rec['group_name']} ({rec['confidence']}% confidence)"
        rec_h = doc.add_heading(rec_title, 3)
        for run in rec_h.runs:
            run.font.size = Pt(12)

        # Rationale
        doc.add_heading('Rationale', 4)
        doc.add_paragraph(rec.get('rationale', ''))

        # Evidence
        if rec.get('evidence'):
            doc.add_heading('Supporting Evidence', 4)
            for ev in rec['evidence']:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.5)
                para.style = 'Quote'
                quote_run = para.add_run(f'"{ev["quote"]}"')
                quote_run.italic = True
                para.add_run(f'\n-- {ev["source_field"]}')

        # Authoritative Source
        doc.add_heading('Authoritative Source', 4)
        provenance = rec.get('provenance', {})

        # Source type
        para = doc.add_paragraph()
        para.add_run('Source: ').bold = True
        para.add_run(provenance.get('source_type', ''))

        # Document (with hyperlink)
        para = doc.add_paragraph()
        para.add_run('Document: ').bold = True
        if provenance.get('url'):
            _add_hyperlink(
                para,
                provenance['url'],
                f"TBS Occupational Group {rec['group_code']} Definition"
            )
        else:
            para.add_run(f"TBS Occupational Group {rec['group_code']} Definition")

        # Inclusions
        if provenance.get('inclusions_referenced'):
            para = doc.add_paragraph()
            para.add_run('Inclusions Referenced: ').bold = True
            para.add_run(', '.join(provenance['inclusions_referenced']))

        # Exclusions
        if provenance.get('exclusions_checked'):
            para = doc.add_paragraph()
            para.add_run('Exclusions Checked: ').bold = True
            para.add_run(', '.join(provenance['exclusions_checked']))

        # Data retrieved
        if provenance.get('scraped_at'):
            para = doc.add_paragraph()
            run = para.add_run(f"Data retrieved: {provenance['scraped_at']}")
            run.font.size = Pt(9)
            run.italic = True

    # Classification Audit Footer
    doc.add_heading('Classification Audit Trail', 3)
    analyzed_at = data.classification_result.get('analyzed_at')
    if analyzed_at and hasattr(analyzed_at, 'strftime'):
        analyzed_str = analyzed_at.strftime('%Y-%m-%d %H:%M UTC')
    else:
        analyzed_str = 'N/A'

    _add_key_value_table(doc, [
        ('Tool', 'JobForge Classification Engine'),
        ('Version', '4.1'),
        ('Generated', analyzed_str),
        ('Data Sources', 'TBS Occupational Group Definitions (2026-01)'),
        ('Compliance', 'TBS Directive 32592 (Automated Decision Making)'),
        ('Constraints', data.classification_result.get('constraints_compliance', ''))
    ])


def generate_docx(data: ExportData) -> bytes:
    """
    Generate Word document from export data using v5.1 section structure.

    Section order:
    1. Title / cover
    2. Position Overview (if AI-generated)
    3. Key Duties and Responsibilities (key_activities with source tags)
    4. Qualifications and Requirements (Skills / Abilities / Knowledge / Core Competencies)
    5. Effort & Physical Demands
    6. Responsibilities
    7. Classification results (if included)
    8. Appendix A: Data Provenance & Compliance
    9. Appendix B: Policy Provenance
    10. Appendix C: Data Quality (DAMA DMBOK)
    11. Additional Job Information (Annex, if present)

    Args:
        data: Complete export data structure

    Returns:
        DOCX bytes ready for response
    """
    from datetime import datetime as _dt

    doc = Document()

    # Page setup
    doc_section = doc.sections[0]
    doc_section.page_height = Inches(11)
    doc_section.page_width = Inches(8.5)
    doc_section.top_margin = Inches(1)
    doc_section.bottom_margin = Inches(1)
    doc_section.left_margin = Inches(0.75)
    doc_section.right_margin = Inches(0.75)

    # Header
    header = doc_section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{data.job_title} ({data.noc_code})"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = GC_PRIMARY

    # Footer
    footer = doc_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Compliant with TBS Directive 32592 | JobForge JD Builder 1.0"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── 1. Title / Cover ──────────────────────────────────────────────────────
    title = doc.add_heading(data.job_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"NOC Code: {data.noc_code}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f"Generated: {data.generated_at.strftime('%Y-%m-%d')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Classification badge in cover (if available)
    if data.include_classification and data.classification_result:
        recs = data.classification_result.get('recommendations', [])
        if recs:
            top = recs[0]
            badge_para = doc.add_paragraph(
                f"Occupational Group: {top['group_code']} — {top['group_name']} "
                f"({top['confidence']}% confidence)"
            )
            badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if badge_para.runs:
                badge_para.runs[0].bold = True

    # ── 2. Position Overview ──────────────────────────────────────────────────
    if data.general_overview:
        h = doc.add_heading("Position Overview", 1)
        for run in h.runs:
            run.font.color.rgb = GC_PRIMARY

        if data.ai_metadata:
            ai_note = doc.add_paragraph()
            ai_label = ai_note.add_run(
                f"AI Generated — Model: {data.ai_metadata.model} · "
                f"Prompt v{data.ai_metadata.prompt_version}"
                + (" · Modified by User" if data.ai_metadata.modified else "")
            )
            ai_label.font.size = Pt(9)
            ai_label.font.color.rgb = AI_DISCLOSURE_BLUE
            ai_label.italic = True

        doc.add_paragraph(data.general_overview)

    # Helper: build element lookup
    elements_by_key = {el.key: el for el in data.jd_elements}

    # Source tag helper (uses module-level SOURCE_TAG_MAP)
    def _src_tag(source_attribute):
        return get_source_tag(source_attribute)

    # ── 3. Key Duties and Responsibilities ───────────────────────────────────
    key_act = elements_by_key.get('key_activities')
    if key_act and key_act.statements:
        h = doc.add_heading("Key Duties and Responsibilities", 1)
        for run in h.runs:
            run.font.color.rgb = GC_PRIMARY
        for stmt in key_act.statements:
            text = stmt.styled_variant.styled_text if stmt.styled_variant else stmt.text
            tag = _src_tag(stmt.source_attribute)
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(text)
            tag_run = para.add_run(f"  {tag}")
            tag_run.font.size = Pt(8)
            tag_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 4. Qualifications and Requirements ───────────────────────────────────
    skills_el = elements_by_key.get('skills')
    cc_sels = [s for s in data.manager_selections if s.jd_element == 'core_competencies']
    if skills_el or cc_sels:
        h = doc.add_heading("Qualifications and Requirements", 1)
        for run in h.runs:
            run.font.color.rgb = GC_PRIMARY

        if skills_el:
            skill_stmts = [s for s in skills_el.statements
                           if s.source_attribute not in ('Abilities', 'Knowledge')]
            if skill_stmts:
                sh = doc.add_heading("Skills", 2)
                for run in sh.runs:
                    run.font.size = Pt(12)
                for stmt in skill_stmts:
                    para = doc.add_paragraph(stmt.text, style='List Bullet')

            ability_stmts = [s for s in skills_el.statements if s.source_attribute == 'Abilities']
            if ability_stmts:
                ah = doc.add_heading("Abilities", 2)
                for run in ah.runs:
                    run.font.size = Pt(12)
                for stmt in ability_stmts:
                    doc.add_paragraph(stmt.text, style='List Bullet')

            know_stmts = [s for s in skills_el.statements if s.source_attribute == 'Knowledge']
            if know_stmts:
                kh = doc.add_heading("Knowledge", 2)
                for run in kh.runs:
                    run.font.size = Pt(12)
                for stmt in know_stmts:
                    doc.add_paragraph(stmt.text, style='List Bullet')

        if cc_sels:
            cch = doc.add_heading("Core Competencies", 2)
            for run in cch.runs:
                run.font.size = Pt(12)
            for sel in cc_sels:
                doc.add_paragraph(sel.text, style='List Bullet')

    # ── 5. Effort & Physical Demands ─────────────────────────────────────────
    effort_el = elements_by_key.get('effort')
    if effort_el and effort_el.statements:
        h = doc.add_heading("Effort & Physical Demands", 1)
        for run in h.runs:
            run.font.color.rgb = GC_PRIMARY
        for stmt in effort_el.statements:
            doc.add_paragraph(stmt.text, style='List Bullet')

    # ── 6. Responsibilities ───────────────────────────────────────────────────
    resp_el = elements_by_key.get('responsibility')
    if resp_el and resp_el.statements:
        h = doc.add_heading("Responsibilities", 1)
        for run in h.runs:
            run.font.color.rgb = GC_PRIMARY
        for stmt in resp_el.statements:
            doc.add_paragraph(stmt.text, style='List Bullet')

    # ── 7. Classification Results ─────────────────────────────────────────────
    _add_classification_section(doc, data)

    # ── 8. Appendix A: Data Provenance & Compliance ───────────────────────────
    doc.add_page_break()
    h = doc.add_heading("Appendix A: Data Provenance & Compliance", 0)

    doc.add_heading("Source Information", 1)
    _add_key_value_table(doc, [
        ("NOC Code", data.noc_code),
        ("Data Source", "JobForge 2.0 / OASIS (Occupational and Skills Information System)"),
        ("Source Authority", "Employment and Social Development Canada (ESDC)"),
        ("Generated By", "JobForge JD Builder 1.0"),
        ("Generation Date", data.generated_at.strftime('%Y-%m-%d %H:%M UTC')),
        ("NOC Version", data.source_metadata.version),
        ("Profile URL", data.source_metadata.profile_url),
        ("Retrieved", data.source_metadata.scraped_at.strftime('%Y-%m-%d %H:%M UTC')),
    ])

    doc.add_heading("DAMA DMBOK Compliance", 1)
    doc.add_paragraph(
        "Data management complies with the GC Data Quality Management Framework (DAMA DMBOK 2.0). "
        "All data sourced from authoritative Government of Canada occupational classification systems "
        "with documented lineage and timestamped retrieval."
    )

    doc.add_heading("DADM Directive Compliance (Section 6.2.7)", 1)
    for sec in data.compliance_sections:
        if sec.section_id == "6.2.7":
            doc.add_paragraph(sec.content.get("description", ""))
            doc.add_paragraph(f"Total selections: {sec.content.get('total_selections', 0)}")
            selections = sec.content.get("selections", [])
            if selections:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                for i, label in enumerate(["JD Element", "Statement", "Source", "Publication Date"]):
                    hdr[i].text = label
                    for para in hdr[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
                for sel in selections:
                    row = table.add_row().cells
                    row[0].text = sel.get("jd_element", "")
                    stmt_text = sel.get("statement", "")
                    row[1].text = stmt_text[:100] + ("…" if len(stmt_text) > 100 else "")
                    row[2].text = sel.get("source_attribute", "")
                    row[3].text = sel.get("publication_date", "")

    doc.add_heading("Directive on Classification Compliance (Section 4.1.2)", 1)
    if data.include_classification and data.classification_result:
        recs = data.classification_result.get('recommendations', [])
        if recs:
            top = recs[0]
            doc.add_paragraph(
                f"Occupational group allocation performed per TBS Directive on Classification "
                f"Section 4.1.2. Recommended group: {top['group_code']} — {top['group_name']} "
                f"({top['confidence']}% confidence). Full rationale recorded in Classification section."
            )
    else:
        doc.add_paragraph(
            "Classification Step 1 not yet performed for this job description. "
            "Proceed to the Classify step to generate an occupational group recommendation "
            "per TBS Directive on Classification Section 4.1.2."
        )

    # ── 9. Appendix B: Policy Provenance ─────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix B: Policy Provenance", 0)

    doc.add_heading("Authoritative Data Sources", 1)
    prov_table = doc.add_table(rows=1, cols=4)
    prov_table.style = 'Table Grid'
    prov_hdr = prov_table.rows[0].cells
    for i, label in enumerate(["Source", "Data Steward", "Publication", "Times Removed"]):
        prov_hdr[i].text = label
        for para in prov_hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    retrieval_date = data.source_metadata.scraped_at.strftime('%Y-%m-%d')
    for row_data in [
        ("2021 NOC", "ESDC", retrieval_date, "0 — Direct"),
        ("OaSIS", "ESDC", retrieval_date, "0 — Direct"),
        ("O*NET SOC (US)", "U.S. Dept. of Labor / ETA", "2024", "1 — Crosswalked"),
    ]:
        r = prov_table.add_row().cells
        for i, val in enumerate(row_data):
            r[i].text = val

    doc.add_heading("Degrees of Separation", 2)
    _add_key_value_table(doc, [
        ("0 — Direct", "Data retrieved directly from authoritative source"),
        ("1 — Crosswalked", "Data mapped from one authoritative taxonomy to another"),
        ("2+ — Derived", "Data computed or inferred from source data"),
    ])

    if data.ai_metadata:
        doc.add_heading("AI-Generated Content", 1)
        _add_key_value_table(doc, [
            ("Content Type", "Position Overview (General Overview)"),
            ("Model", data.ai_metadata.model),
            ("Generation Timestamp", data.ai_metadata.timestamp.strftime('%Y-%m-%d %H:%M UTC')),
            ("Prompt Version", data.ai_metadata.prompt_version),
            ("Input Statements", str(len(data.ai_metadata.input_statement_ids))),
            ("Modified by User", "Yes" if data.ai_metadata.modified else "No"),
        ])

    # ── 10. Appendix C: Data Quality ─────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Appendix C: Data Quality (DAMA DMBOK)", 0)
    for sec in data.compliance_sections:
        if sec.section_id == "6.3.5":
            c = sec.content
            _add_key_value_table(doc, [
                ("Accuracy", c.get('accurate', {}).get('how_met', '')),
                ("Completeness",
                 "All selected JD elements retrieved from authoritative NOC/OaSIS sources "
                 "with no truncation"),
                ("Consistency",
                 "NOC 2021 v1.0 used as single authoritative taxonomy; "
                 "crosswalks documented in provenance chain"),
                ("Timeliness", c.get('up_to_date', {}).get('how_met', '')),
            ])

    # ── 11. Additional Job Information (Annex) ────────────────────────────────
    _add_annex_section(doc, data)

    # Write to bytes
    with BytesIO() as buffer:
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()


def _add_key_value_table(doc, items):
    """Helper to add a simple 2-column key-value table."""
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'

    for i, (key, value) in enumerate(items):
        cells = table.rows[i].cells
        cells[0].text = key
        cells[1].text = str(value)

        # Make key column bold
        for para in cells[0].paragraphs:
            for run in para.runs:
                run.bold = True


# Text color constant for light gray
TEXT_LIGHT = RGBColor(0x66, 0x66, 0x66)


def _add_styled_statement(doc, statement):
    """Add statement with styled content and original below.

    Args:
        doc: Document object
        statement: StatementExport with styled_variant
    """
    styled = statement.styled_variant

    # Primary: Styled text as bullet point
    para = doc.add_paragraph()
    para.style = 'List Bullet'
    text_run = para.add_run(styled.styled_text)

    # Confidence dot (only for AI_STYLED, not ORIGINAL_NOC fallback)
    if styled.content_type == StyleContentType.AI_STYLED:
        dot_run = para.add_run(" \u25cf")  # Unicode filled circle
        dot_run.font.color.rgb = CONFIDENCE_COLORS.get(
            _get_confidence_level(styled.confidence_score),
            CONFIDENCE_COLORS["low"]
        )
        dot_run.font.size = Pt(8)

    # Original NOC text (only for AI_STYLED, not ORIGINAL_NOC fallback)
    if styled.content_type == StyleContentType.AI_STYLED:
        orig_para = doc.add_paragraph()
        orig_para.paragraph_format.left_indent = Inches(0.5)

        label_run = orig_para.add_run("Original NOC: ")
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        label_run.bold = True

        text_run = orig_para.add_run(styled.original_noc_text)
        text_run.font.size = Pt(9)
        text_run.font.color.rgb = TEXT_LIGHT
        text_run.italic = True

    # AI disclosure label (for any styled variant)
    label_para = doc.add_paragraph()
    label_para.paragraph_format.left_indent = Inches(0.25)
    label_run = label_para.add_run(styled.disclosure_label)
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = AI_DISCLOSURE_BLUE
    label_run.italic = True


def _add_annex_section(doc, data):
    """Add Annex section with unused NOC attributes.

    Renders the Annex after compliance appendix with all 4 categories:
    - Job Requirements (paragraph format)
    - Career Mobility (grouped list)
    - Interests (Holland Codes) (bullet list)
    - Personal Suitability (Placement Criteria) (bullet list)
    """
    if not data.annex_data or not data.annex_data.sections:
        return  # No annex data to render

    # Page break before Annex
    doc.add_page_break()

    # Main heading (Heading 1 for navigation pane)
    main_heading = doc.add_heading('Additional Job Information', 1)
    for run in main_heading.runs:
        run.font.color.rgb = GC_PRIMARY

    # Intro paragraph
    intro = doc.add_paragraph(
        'Reference information from the National Occupational Classification '
        'that may be useful for recruitment, career development, and job analysis.'
    )
    if intro.runs:
        intro.runs[0].italic = True

    # Render each category section
    for section in data.annex_data.sections:
        # Category heading (Heading 2 for hierarchy)
        category_heading = doc.add_heading(section.title, 2)
        for run in category_heading.runs:
            run.font.color.rgb = GC_PRIMARY
            run.font.size = Pt(12)
        category_heading.paragraph_format.space_after = Pt(9)

        # Handle empty sections
        if not section.items:
            empty_para = doc.add_paragraph('No additional information available.')
            if empty_para.runs:
                empty_para.runs[0].italic = True
                empty_para.runs[0].font.color.rgb = TEXT_LIGHT
            continue

        # Render based on format type
        if section.format_type == 'paragraph':
            # Job Requirements: paragraphs
            for item in section.items:
                doc.add_paragraph(item)

        elif section.format_type == 'grouped_list':
            # Career Mobility: grouped list with Entry/Advancement headers
            for item in section.items:
                if item.endswith(':'):
                    # Group header (Entry Paths:, Advancement Paths:)
                    para = doc.add_paragraph()
                    run = para.add_run(item)
                    run.bold = True
                else:
                    # Mobility item (indented with dash)
                    para = doc.add_paragraph(item)
                    para.paragraph_format.left_indent = Inches(0.25)

        else:  # 'list'
            # Interests, Personal Suitability: bullet list
            for item in section.items:
                para = doc.add_paragraph(item, style='List Bullet')
                para.paragraph_format.space_after = Pt(6)

    # Source attribution at end of Annex
    attr_para = doc.add_paragraph()
    attr_para.paragraph_format.space_before = Pt(24)
    attr_run = attr_para.add_run(
        f"Source: NOC {data.annex_data.source_noc_code}, "
        f"retrieved {data.annex_data.retrieved_at.strftime('%Y-%m-%d')}"
    )
    attr_run.font.size = Pt(9)
    attr_run.font.color.rgb = TEXT_LIGHT
    attr_run.italic = True
